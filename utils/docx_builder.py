"""워드(.docx) 문서 빌더 — 교육·평가 시험지 / SOP 초안.

python-docx 사용. 모든 함수는 bytes(docx 파일)를 반환해
Streamlit st.download_button 으로 바로 내려받을 수 있다.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm


def _title(doc: Document, text: str, size: int = 18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _info_table(doc: Document, rows: List[tuple]):
    """2열(항목/값) 정보 표."""
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k1, v1, k2, v2) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = k1
        cells[1].text = v1
        cells[2].text = k2
        cells[3].text = v2
        for c in (cells[0], cells[2]):
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True


# ---------------------------------------------------------------- 시험지

def build_exam_docx(
    questions,                       # List[exam_engine.Question]
    *,
    company: str = "",
    exam_title: str = "KGSP 품질관리 교육 평가",
    exam_date: str = "",
    trainer: str = "",
    pass_score: int = 80,
    include_answer_key: bool = True,
    include_training_log: bool = True,
    attendee_rows: int = 10,
) -> bytes:
    """시험지(+정답·해설지, +교육일지) 한 파일로 생성 → bytes."""
    exam_date = exam_date or datetime.now().strftime("%Y-%m-%d")
    doc = Document()

    # ---- 1) 표지 + 응시자 정보
    _title(doc, exam_title)
    if company:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(company)
    doc.add_paragraph()
    _info_table(doc, [
        ("교육일자", exam_date, "소속/부서", ""),
        ("성명", "", "점수", f"        /100  (합격 {pass_score}점)"),
    ])
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(
        f"※ 총 {len(questions)}문항. OX 문항은 괄호에 O 또는 X 를, "
        "객관식은 번호(①~④)를 기입하세요."
    ).font.size = Pt(9)
    doc.add_paragraph()

    # ---- 2) 문항
    per_q = round(100 / len(questions), 1) if questions else 0
    for q in questions:
        head = doc.add_paragraph()
        run = head.add_run(f"{q.number}. {q.question}")
        run.bold = False
        if q.qtype == "ox":
            doc.add_paragraph("      답: (        )")
        else:
            for i, opt in enumerate(q.options):
                doc.add_paragraph(f"      {['①','②','③','④'][i]} {opt}")
            doc.add_paragraph("      답: (        )")
    p = doc.add_paragraph()
    p.add_run(f"(문항당 {per_q}점)").font.size = Pt(9)

    # ---- 3) 정답·해설 (별지)
    if include_answer_key:
        doc.add_page_break()
        _title(doc, "정답 및 해설 (관리자 보관용)", 15)
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "문항"
        hdr[1].text = "정답"
        hdr[2].text = "해설 / 근거"
        for q in questions:
            row = table.add_row().cells
            row[0].text = str(q.number)
            row[1].text = q.answer
            expl = q.explanation or ""
            if q.source and q.source not in expl:
                expl = (expl + f" (근거: {q.source})").strip()
            row[2].text = expl

    # ---- 4) 교육일지 (별지)
    if include_training_log:
        doc.add_page_break()
        _title(doc, "교육 실시 기록 (교육일지)", 15)
        doc.add_paragraph()
        _info_table(doc, [
            ("교육일자", exam_date, "교육방법", "집합교육 □  자체학습 □"),
            ("교육명", exam_title, "교육시간", ""),
            ("강사/주관", trainer, "평가방법", f"필기평가 (합격 {pass_score}점 이상)"),
        ])
        doc.add_paragraph()
        doc.add_paragraph("참석자 명단 및 평가 결과")
        table = doc.add_table(rows=attendee_rows + 1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["번호", "소속", "성명", "점수", "서명"]):
            hdr[i].text = h
        for i in range(1, attendee_rows + 1):
            table.rows[i].cells[0].text = str(i)
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig.add_run("관리약사(품질책임자) 확인:  ______________  (인)")

    foot = doc.add_paragraph()
    foot.add_run(
        "본 평가지는 자동 생성된 교육 보조 자료입니다 — 사용 전 QA 검토 필요."
    ).font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------- SOP 초안

def build_sop_docx(
    sections: List[tuple],           # [(제목, 본문), ...]
    *,
    sop_title: str = "표준작업절차서(SOP)",
    company: str = "",
    doc_no: str = "",
    version: str = "1.0",
    template_bytes: Optional[bytes] = None,
) -> bytes:
    """SOP 초안 docx 생성.

    template_bytes 가 있으면 그 문서를 베이스로 사용 — 머리글/바닥글·스타일·
    페이지 설정이 그대로 유지되고, 본문만 새 내용으로 교체된다.
    없으면 표준 양식(문서정보 표 + 개정이력 + 조 단위 본문)으로 생성.
    """
    if template_bytes:
        doc = Document(io.BytesIO(template_bytes))
        # 본문 비우기 (머리글/바닥글/스타일은 문서에 남는다)
        body = doc.element.body
        for el in list(body):
            if el.tag.endswith("}sectPr"):  # 페이지 설정은 유지
                continue
            body.remove(el)
    else:
        doc = Document()

    _title(doc, sop_title)
    doc.add_paragraph()
    _info_table(doc, [
        ("문서번호", doc_no, "개정번호", version),
        ("제정/개정일", datetime.now().strftime("%Y-%m-%d"), "작성", ""),
        ("회사명", company, "승인(품질책임자)", ""),
    ])
    doc.add_paragraph()

    # 개정 이력
    doc.add_paragraph("개정 이력")
    hist = doc.add_table(rows=2, cols=4)
    hist.style = "Table Grid"
    for i, h in enumerate(["개정번호", "개정일", "개정 내용", "승인"]):
        hist.rows[0].cells[i].text = h
    hist.rows[1].cells[0].text = version
    hist.rows[1].cells[1].text = datetime.now().strftime("%Y-%m-%d")
    hist.rows[1].cells[2].text = "최초 제정(자동 생성 초안)"
    doc.add_paragraph()

    # 본문 (조 단위)
    for heading, body_text in sections:
        h = doc.add_paragraph()
        h.add_run(heading).bold = True
        for line in (body_text or "").split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        doc.add_paragraph()

    foot = doc.add_paragraph()
    foot.add_run(
        "본 문서는 자동 생성된 초안입니다 — 시행 전 반드시 품질책임자(관리약사) 검토·승인 필요."
    ).font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
