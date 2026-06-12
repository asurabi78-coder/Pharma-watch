"""업로드 문서 → 텍스트 추출 유틸 — docx / pdf / txt / md 지원.

Streamlit file_uploader 의 UploadedFile(또는 bytes)을 받아 본문 텍스트를 꺼낸다.
LLM 미사용(결정론적). 라이브러리가 없으면 친절한 안내 문자열을 에러로 반환.

사용:
    text, err = extract_text(uploaded.name, uploaded.getvalue())
    if err: st.error(err)
"""
from __future__ import annotations

import io
from typing import Optional, Tuple

SUPPORTED_EXTS = (".docx", ".pdf", ".txt", ".md")


def _from_docx(data: bytes) -> str:
    from docx import Document  # python-docx
    doc = Document(io.BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # 표 안의 텍스트도 포함 (SOP 는 표 양식이 많다)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for pg in reader.pages:
        try:
            t = (pg.extract_text() or "").strip()
        except Exception:
            t = ""
        if t:
            pages.append(t)
    return "\n".join(pages)


def _from_text(data: bytes) -> str:
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, data: bytes) -> Tuple[str, Optional[str]]:
    """(텍스트, 에러메시지) — 성공 시 에러는 None.

    스캔본 PDF(이미지)는 텍스트가 안 나올 수 있다 → 안내 에러 반환.
    """
    name = (filename or "").lower()
    try:
        if name.endswith(".docx"):
            text = _from_docx(data)
        elif name.endswith(".pdf"):
            text = _from_pdf(data)
        elif name.endswith((".txt", ".md")):
            text = _from_text(data)
        elif name.endswith(".doc"):
            return "", ("구형 .doc 형식은 지원하지 않습니다 — Word 에서 "
                        "'다른 이름으로 저장 → .docx' 후 다시 업로드하세요.")
        elif name.endswith((".hwp", ".hwpx")):
            return "", ("한글(.hwp) 파일은 지원하지 않습니다 — 한글에서 "
                        "'다른 이름으로 저장 → DOCX 또는 PDF' 후 다시 업로드하세요.")
        else:
            return "", f"지원하지 않는 형식입니다 ({filename}) — docx/pdf/txt/md 만 가능합니다."
    except ImportError as e:
        missing = "python-docx" if "docx" in str(e).lower() else "pypdf"
        return "", (f"서버에 {missing} 라이브러리가 없습니다 — "
                    f"`pip install {missing}` 후 다시 시도하세요.")
    except Exception as e:  # noqa: BLE001
        return "", f"파일을 읽지 못했습니다: {type(e).__name__}: {e}"

    text = (text or "").strip()
    if not text:
        return "", ("문서에서 텍스트를 찾지 못했습니다 — 스캔본(이미지) PDF 이거나 "
                    "빈 문서일 수 있습니다. 텍스트가 포함된 파일로 다시 시도하세요.")
    return text, None
